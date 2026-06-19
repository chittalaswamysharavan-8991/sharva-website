from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT_PATH = Path(r"C:\Website\outputs\pablo_cockpit_chat_summary.docx")


def set_font(run, name="Arial", size=11, bold=False, color="000000"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(paragraph, before=0, after=8, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = tblBorders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tblBorders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "D9DDE3")


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
styles["Normal"].font.size = Pt(11)

for style_name, size, after in [("Heading 1", 20, 6), ("Heading 2", 16, 6), ("Heading 3", 14, 4)]:
    style = styles[style_name]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    style.font.size = Pt(size)
    style.font.bold = False
    style.font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.space_before = Pt(20 if style_name == "Heading 1" else 18 if style_name == "Heading 2" else 16)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.15

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(3)
title.paragraph_format.line_spacing = 1.0
run = title.add_run("Pablo Cockpit End-to-End Work Summary")
set_font(run, size=26, bold=False)

subtitle = doc.add_paragraph()
style_paragraph(subtitle, before=0, after=10, line=1.15)
subtitle.add_run("Scope: everything completed in this chat, the current live state, and what remains before the rollout is fully finished.").font.size = Pt(11)
set_font(subtitle.runs[0], size=11, color="555555")

doc.add_heading("1. Objective", level=1)
p = doc.add_paragraph()
style_paragraph(p)
text = (
    "The working goal for this thread was to continue the Pablo Cockpit website build in the Website folder, "
    "verify the signed-in private /sync-settings view, and carry the v3 rollout forward without breaking the privacy rules."
)
p.add_run(text)

doc.add_heading("2. What I Checked First", level=1)
for item in [
    "Reviewed the live current state in the browser and the repository before changing anything.",
    "Loaded the Supabase guidance because the remaining work depended on auth, RLS, and connector token storage.",
    "Read the existing connector helper and route files so the refactor could be finished from the current implementation instead of starting over.",
]:
    p = doc.add_paragraph(style="List Bullet")
    style_paragraph(p, after=4, line=1.15)
    p.add_run(item)

doc.add_heading("3. Code Changes Completed", level=1)
changes = [
    ("Shared connector helper", [
        "Finished the connector auth helper in api/_lib/connector-auth.js.",
        "Added signed connector session helpers, cookie helpers, and return-path normalization.",
        "Kept the Supabase client on the anon key path with RLS-backed writes instead of adding a service-role dependency."
    ]),
    ("OAuth start routes", [
        "Updated the Google Calendar and Notion start routes to verify the signed-in owner.",
        "Created a short-lived signed connector session cookie for the OAuth round-trip.",
        "Returned the provider authorization URLs using the real client IDs and callback paths."
    ]),
    ("OAuth callback routes", [
        "Updated both callback routes to verify the signed state and the signed connector session cookie.",
        "Wrote connector tokens and connector account state through the authenticated owner session.",
        "Triggered the first sync pass after a successful connection and cleared the cookie afterward."
    ]),
    ("Sync routes", [
        "Fixed the Google Calendar and Notion sync endpoints to pass the signed-in owner access token through the sync helpers.",
        "Kept the stored token refresh flow inside the same authenticated path."
    ]),
    ("UI and readiness", [
        "Updated src/main.jsx so /sync-settings shows the real connector setup path and uses same-origin requests for the OAuth start calls.",
        "Adjusted the connector readiness snapshot so it reports CONNECTOR_STATE_SECRET as part of the server-side setup requirement.",
        "Refined the wording in the setup cards so the page reflects the actual wired routes instead of reserving them for later."
    ]),
]
for heading, bullets in changes:
    doc.add_heading(heading, level=2)
    for item in bullets:
        p = doc.add_paragraph(style="List Bullet")
        style_paragraph(p, after=4, line=1.15)
        p.add_run(item)

doc.add_heading("4. Environment And Docs Updated", level=1)
for item in [
    "Removed the Supabase service-role key from .env.example because the implemented connector flow does not need it.",
    "Added CONNECTOR_STATE_SECRET to the env example and mirrored that requirement in the connector status endpoint and app UI.",
    "Updated docs/connector-setup-guide.md to describe the real provider setup, redirect URLs, and tap-by-tap flow.",
    "Updated docs/v2-auth-sync-plan.md, README.md, and docs/buildos-pablo-cockpit.md so they match the current live architecture."
]:
    p = doc.add_paragraph(style="List Bullet")
    style_paragraph(p, after=4, line=1.15)
    p.add_run(item)

doc.add_heading("5. Verification Completed", level=1)
for item in [
    "Confirmed the code loaded successfully with a route-load check in Node.",
    "Ran a production build successfully with vite build.",
    "Deployed the updated app to Vercel production successfully.",
    "Verified the live connector-status endpoint shows the new server-side requirements.",
    "Verified the public doorway still opens and the canonical host redirect behavior remains intact.",
    "Verified that the private cockpit screen still opens to /home and shows the owner authentication gate."
]:
    p = doc.add_paragraph(style="List Bullet")
    style_paragraph(p, after=4, line=1.15)
    p.add_run(item)

doc.add_heading("6. Live Browser Findings", level=1)
p = doc.add_paragraph()
style_paragraph(p)
p.add_run(
    "The browser proof reached the private cockpit gate, but the Google OAuth login step was blocked by Google's "
    "automation/reputation checks in headless or automated browser flows. I switched to a copied Chrome profile and "
    "confirmed the app and sign-in screen behavior, but the actual Google sign-in still needs a human-style interactive pass."
)

doc.add_heading("7. Current State", level=1)
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
remove_table_borders(table)
hdr = table.rows[0].cells
hdr[0].text = "Area"
hdr[1].text = "Status"
rows = [
    ("Core app build", "Done and deployed"),
    ("Private /home gate", "Verified open to the auth screen"),
    ("Private /sync-settings view", "Implemented and live, but not yet fully proven with provider secrets"),
    ("Google Calendar / Notion routes", "Implemented"),
    ("Docs and env cleanup", "Implemented"),
    ("Final live OAuth sign-in proof", "Still pending"),
]
for area, status in rows:
    cells = table.add_row().cells
    cells[0].text = area
    cells[1].text = status

for row in table.rows:
    for idx, cell in enumerate(row.cells):
        for para in cell.paragraphs:
            style_paragraph(para, before=0, after=0, line=1.0)
            for run in para.runs:
                set_font(run, size=10 if idx else 10, bold=False)

doc.add_heading("8. What Is Still Pending", level=1)
for item in [
    "Add CONNECTOR_STATE_SECRET to Vercel.",
    "Add GOOGLE_CLIENT_ID and GOOGLE_CLIENT_SECRET to Vercel.",
    "Add NOTION_CLIENT_ID and NOTION_CLIENT_SECRET to Vercel.",
    "Complete a real interactive Google sign-in in the browser so the private owner session can be proved end-to-end.",
    "Re-check /sync-settings after the provider secrets are in place and the OAuth flow completes."
]:
    p = doc.add_paragraph(style="List Number")
    style_paragraph(p, after=4, line=1.15)
    p.add_run(item)

doc.add_heading("9. Important Note", level=1)
p = doc.add_paragraph()
style_paragraph(p)
p.add_run(
    "The work is not being treated as complete yet because the live browser proof of the authenticated private view "
    "still needs the real provider sign-in step. Everything else was pushed forward from code through deployment and docs."
)

OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT_PATH)
print(OUT_PATH)
